import os
import google.generativeai as genai
from google.generativeai import GenerativeModel
from dotenv import load_dotenv
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document  # Importing for Word generation
import textwrap

# Load the API key from the .env file
load_dotenv()
genai.configure(api_key=os.getenv('GEMINI_API_KEY'))

# Generation settings to control the model's output
generation_config = {
    'temperature': 1,
    'top_p': 0.95,
    'top_k': 64,
    'max_output_tokens': 8192,
    'response_mime_type': 'text/plain'
}

# Safety settings to control the model's content filtering
safety_settings = [
    {'category': 'HARM_CATEGORY_HARASSMENT', 'threshold': 'BLOCK_NONE'},
    {'category': 'HARM_CATEGORY_HATE_SPEECH', 'threshold': 'BLOCK_NONE'},
    {'category': 'HARM_CATEGORY_SEXUALLY_EXPLICIT', 'threshold': 'BLOCK_NONE'},
    {'category': 'HARM_CATEGORY_DANGEROUS_CONTENT', 'threshold': 'BLOCK_NONE'}
]

# Load the instructions for the model
with open('app/instructions.md', 'r') as file:
    gemini_instructions = file.read()

# Initialize the Gemini model with custom settings and instructions
llm = GenerativeModel(
    model_name='gemini-1.5-flash-latest',
    generation_config=generation_config,
    safety_settings=safety_settings,
    system_instruction=gemini_instructions
)

# Start a chat session with the model
chat_session = llm.start_chat(history=[])

def generate_summary(file_data):
    user_message = f'File Data: {file_data}'
    bot_response = chat_session.send_message(user_message)
    return bot_response.text

def generate_pdf_from_response(response_text, output_file='output.pdf'):
    """Generates a PDF file with the bot's response, handling line wrapping and simple formatting."""
    c = canvas.Canvas(output_file, pagesize=letter)
    width, height = letter
    margin = inch  # 1-inch margins
    text_width = width - 2 * margin
    text_height = height - margin

    # Start writing from the top of the page
    text_object = c.beginText(margin, text_height)
    text_object.setFont("Helvetica", 12)

    # Wrap lines to fit within the page width
    wrapper = textwrap.TextWrapper(width=int(text_width / 6))  # Adjust width to match PDF text width

    # Split the response into lines and format based on simple rules
    lines = response_text.split('\n')
    for line in lines:
        if line.startswith('# '):  # Treat as a heading
            text_object.setFont("Helvetica-Bold", 14)
            wrapped_text = wrapper.wrap(line[2:])  # Remove '# ' and wrap text
        elif line.startswith('**') and line.endswith('**'):  # Treat as bold text
            text_object.setFont("Helvetica-Bold", 12)
            wrapped_text = wrapper.wrap(line[2:-2])  # Remove '**' and wrap text
        else:  # Regular text
            text_object.setFont("Helvetica", 12)
            wrapped_text = wrapper.wrap(line)

        # Add the wrapped text to the PDF line by line
        for wrapped_line in wrapped_text:
            text_object.textLine(wrapped_line)
    
    # Write the text to the PDF
    c.drawText(text_object)
    c.save()

def generate_word_from_response(response_text, output_file='output.docx'):
    """Generates a Word (.docx) file with the bot's response, handling simple formatting."""
    doc = Document()

    # Split the response into lines and format based on simple rules
    lines = response_text.split('\n')
    for line in lines:
        if line.startswith('# '):  # Treat as a heading
            doc.add_heading(line[2:], level=1)  # Remove '# ' and add as heading
        elif line.startswith('**') and line.endswith('**'):  # Treat as bold text
            doc.add_paragraph(line[2:-2], style='BodyText').bold = True  # Remove '**' and make bold
        else:  # Regular text
            doc.add_paragraph(line)

    # Save the Word document
    doc.save(output_file)

# Example usage
if __name__ == '__main__':
    file_data = "give me 5000 word of on the topic of ai and its benefitd elaborate"
    
    # Get the bot's summary
    summary = generate_summary(file_data)
    
    # Generate a PDF from the bot's response
    generate_pdf_from_response(summary, output_file='bot_response.pdf')

    generate_word_from_response(summary, output_file='bot_response.docx')

